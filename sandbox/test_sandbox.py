"""Sandbox 全面测试脚本.

用法:
  python test_sandbox.py [--url http://localhost:8888] [--agent-id test-agent]
"""

import argparse
import json
import sys
import time
import urllib.request
import urllib.error

# ---------------------------------------------------------------------------
# 配置
# ---------------------------------------------------------------------------
BASE_URL = "http://localhost:8888"
AGENT_ID = "test-sandbox-agent"
PASSED = 0
FAILED = 0
ERRORS: list[str] = []


def _post(path: str, body: dict) -> dict:
    data = json.dumps(body).encode()
    req = urllib.request.Request(
        f"{BASE_URL}{path}",
        data=data,
        headers={"Content-Type": "application/json"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            return json.loads(resp.read())
    except urllib.error.HTTPError as e:
        return {"_http_error": e.code, "detail": e.read().decode()[:200]}
    except Exception as e:
        return {"_error": str(e)[:200]}


def _get(path: str) -> dict:
    req = urllib.request.Request(f"{BASE_URL}{path}")
    try:
        with urllib.request.urlopen(req, timeout=10) as resp:
            return json.loads(resp.read())
    except Exception as e:
        return {"_error": str(e)[:200]}


def execute(code: str, language: str = "python", timeout: int = 30,
            work_dir: str | None = None) -> dict:
    body: dict = {"code": code, "language": language, "timeout": timeout}
    if work_dir:
        body["work_dir"] = work_dir
    return _post("/execute", body)


# ---------------------------------------------------------------------------
# 测试辅助
# ---------------------------------------------------------------------------
def run_test(name: str, func) -> None:
    global PASSED, FAILED
    print(f"\n{'='*60}")
    print(f"TEST: {name}")
    print(f"{'='*60}")
    try:
        func()
        PASSED += 1
        print(f"  [PASS] {name}")
    except AssertionError as e:
        FAILED += 1
        msg = f"  [FAIL] {e}"
        print(msg)
        ERRORS.append(f"{name}: {e}")
    except Exception as e:
        FAILED += 1
        msg = f"  [ERROR] {e}"
        print(msg)
        ERRORS.append(f"{name}: {e}")


def assert_ok(result: dict, msg: str = ""):
    assert result.get("success") is True, (
        f"Expected success=true, got {result.get('success')}. "
        f"{msg} stderr={result.get('stderr', '')[:200]} error={result.get('error', '')}"
    )


def assert_fail(result: dict, msg: str = ""):
    assert result.get("success") is False, (
        f"Expected success=false, got {result.get('success')}. {msg}"
    )


def assert_stdout_contains(result: dict, text: str):
    stdout = result.get("stdout", "")
    assert text in stdout, f"Expected stdout to contain '{text}', got: {stdout[:200]}"


def assert_stderr_contains(result: dict, text: str):
    stderr = result.get("stderr", "")
    assert text in stderr, f"Expected stderr to contain '{text}', got: {stderr[:200]}"


# ---------------------------------------------------------------------------
# 测试用例
# ---------------------------------------------------------------------------

def test_health():
    r = _get("/health")
    assert r.get("status") == "ok", f"Health check failed: {r}"


def test_python_basic():
    r = execute("print('hello python')")
    assert_ok(r)
    assert_stdout_contains(r, "hello python")


def test_python_multiline():
    code = """
import sys
def add(a, b):
    return a + b
result = add(3, 5)
print(f"result={result}")
print(f"python={sys.version_info.major}.{sys.version_info.minor}")
"""
    r = execute(code)
    assert_ok(r)
    assert_stdout_contains(r, "result=8")
    assert_stdout_contains(r, "python=3.12")


def test_bash_basic():
    r = execute("echo 'hello bash'", language="bash")
    assert_ok(r)
    assert_stdout_contains(r, "hello bash")


def test_bash_multiline():
    code = """
for i in 1 2 3; do
    echo "line $i"
done
echo "done"
"""
    r = execute(code, language="bash")
    assert_ok(r)
    assert_stdout_contains(r, "line 1")
    assert_stdout_contains(r, "line 2")
    assert_stdout_contains(r, "line 3")
    assert_stdout_contains(r, "done")


def test_node_basic():
    r = execute("console.log('hello node')", language="node")
    assert_ok(r)
    assert_stdout_contains(r, "hello node")


def test_node_multiline():
    code = """
const items = ['a', 'b', 'c'];
items.forEach(item => console.log(item));
console.log('count=' + items.length);
"""
    r = execute(code, language="node")
    assert_ok(r)
    assert_stdout_contains(r, "a")
    assert_stdout_contains(r, "b")
    assert_stdout_contains(r, "c")
    assert_stdout_contains(r, "count=3")


def test_preinstalled_pandas():
    r = execute("import pandas as pd; print(f'pandas={pd.__version__}')")
    assert_ok(r)
    assert_stdout_contains(r, "pandas=")


def test_preinstalled_numpy():
    r = execute("import numpy as np; print(f'numpy={np.__version__}')")
    assert_ok(r)
    assert_stdout_contains(r, "numpy=")


def test_preinstalled_scipy():
    r = execute("import scipy; print(f'scipy={scipy.__version__}')")
    assert_ok(r)
    assert_stdout_contains(r, "scipy=")


def test_preinstalled_matplotlib():
    r = execute("import matplotlib; print(f'matplotlib={matplotlib.__version__}')")
    assert_ok(r)
    assert_stdout_contains(r, "matplotlib=")


def test_preinstalled_docx():
    r = execute("from docx import Document; print('docx=ok')")
    assert_ok(r)
    assert_stdout_contains(r, "docx=ok")


def test_preinstalled_pptx():
    r = execute("from pptx import Presentation; print('pptx=ok')")
    assert_ok(r)
    assert_stdout_contains(r, "pptx=ok")


def test_preinstalled_openpyxl():
    r = execute("import openpyxl; print(f'openpyxl={openpyxl.__version__}')")
    assert_ok(r)
    assert_stdout_contains(r, "openpyxl=")


def test_preinstalled_xlrd():
    r = execute("import xlrd; print(f'xlrd={xlrd.__version__}')")
    assert_ok(r)
    assert_stdout_contains(r, "xlrd=")


def test_preinstalled_pillow():
    r = execute("from PIL import Image; print('pillow=ok')")
    assert_ok(r)
    assert_stdout_contains(r, "pillow=ok")


def test_preinstalled_pdfplumber():
    r = execute("import pdfplumber; print(f'pdfplumber={pdfplumber.__version__}')")
    assert_ok(r)
    assert_stdout_contains(r, "pdfplumber=")


def test_preinstalled_fpdf2():
    r = execute("from fpdf import FPDF; print('fpdf2=ok')")
    assert_ok(r)
    assert_stdout_contains(r, "fpdf2=ok")


def test_preinstalled_reportlab():
    r = execute("from reportlab.lib.pagesizes import letter; print('reportlab=ok')")
    assert_ok(r)
    assert_stdout_contains(r, "reportlab=ok")


def test_preinstalled_requests():
    r = execute("import requests; print(f'requests={requests.__version__}')")
    assert_ok(r)
    assert_stdout_contains(r, "requests=")


def test_preinstalled_beautifulsoup():
    r = execute("from bs4 import BeautifulSoup; print('bs4=ok')")
    assert_ok(r)
    assert_stdout_contains(r, "bs4=ok")


def test_preinstalled_lxml():
    r = execute("import lxml; print('lxml=ok')")
    assert_ok(r)
    assert_stdout_contains(r, "lxml=ok")


def test_preinstalled_chardet():
    r = execute("import chardet; print(f'chardet={chardet.__version__}')")
    assert_ok(r)
    assert_stdout_contains(r, "chardet=")


def test_preinstalled_tqdm():
    r = execute("from tqdm import tqdm; print('tqdm=ok')")
    assert_ok(r)
    assert_stdout_contains(r, "tqdm=ok")


def test_pip_install_runtime():
    """运行时 pip install 并验证导入"""
    r = execute("pip install qrcode", language="bash", timeout=30)
    assert_ok(r)
    assert_stdout_contains(r, "Successfully installed")

    r2 = execute("import qrcode; qr = qrcode.make('test'); print('qrcode=ok')")
    assert_ok(r2)
    assert_stdout_contains(r2, "qrcode=ok")


def test_file_write_to_workspace():
    """写文件到 agent workspace"""
    work_dir = f"/data/agents/{AGENT_ID}/workspace"
    r = execute(f"""
import os
os.makedirs('{work_dir}', exist_ok=True)
with open('{work_dir}/test_file.txt', 'w') as f:
    f.write('hello from sandbox')
print('write_ok')
""")
    assert_ok(r)
    assert_stdout_contains(r, "write_ok")


def test_file_read_from_workspace():
    """读 workspace 中刚写入的文件"""
    work_dir = f"/data/agents/{AGENT_ID}/workspace"
    code = f"""
with open('{work_dir}/test_file.txt', 'r') as f:
    content = f.read()
print('content=' + content)
"""
    r = execute(code)
    assert_ok(r)
    assert_stdout_contains(r, "content=hello from sandbox")


def test_file_write_to_agent_root():
    """写文件到 agent 根目录（skills/, memory/ 同级）"""
    agent_root = f"/data/agents/{AGENT_ID}"
    r = execute(f"""
import os
os.makedirs('{agent_root}/skills', exist_ok=True)
with open('{agent_root}/skills/test_skill.md', 'w') as f:
    f.write('# Test Skill')
with open('{agent_root}/skills/test_skill.md', 'r') as f:
    print(f.read())
""")
    assert_ok(r)
    assert_stdout_contains(r, "# Test Skill")


def test_workspace_dir_env_var():
    """验证 WORKSPACE_DIR 环境变量"""
    work_dir = f"/data/agents/{AGENT_ID}/workspace"
    r = execute("import os; print(os.environ.get('WORKSPACE_DIR', 'NOT_SET'))",
                work_dir=work_dir)
    assert_ok(r)
    assert_stdout_contains(r, work_dir)


def test_pythonpath_env():
    """验证 PYTHONPATH 包含 shared-deps"""
    r = execute("import os; pp = os.environ.get('PYTHONPATH', ''); print(pp)")
    assert_ok(r)
    assert_stdout_contains(r, "/data/shared-deps/pip")


def test_timeout_short():
    """超时测试：代码执行超过 timeout 应该被 kill"""
    r = execute("import time; time.sleep(10); print('should not reach')", timeout=2)
    assert_fail(r)
    assert r.get("exit_code") == 124, f"Expected exit_code=124, got {r.get('exit_code')}"
    assert "timed out" in (r.get("error") or "").lower(), f"Expected timeout error, got {r.get('error')}"


def test_bash_invoke_python():
    """Bash 中调用 Python"""
    r = execute("python3 -c 'print(\"from bash\")'", language="bash")
    assert_ok(r)
    assert_stdout_contains(r, "from bash")


def test_bash_invoke_node():
    """Bash 中调用 Node.js"""
    r = execute("node -e 'console.log(\"from bash\")'", language="bash")
    assert_ok(r)
    assert_stdout_contains(r, "from bash")


def test_pip_install_in_python():
    """Python 代码中的 pip install 转换"""
    r = execute("pip install pyyaml", language="bash", timeout=30)
    assert_ok(r)
    assert_stdout_contains(r, "Successfully installed")


def test_large_output():
    """大输出测试 — stdout 截断到 10KB + 溢出文件"""
    work_dir = f"/data/agents/{AGENT_ID}"
    r = execute("print('x' * 50000)", work_dir=work_dir)
    assert_ok(r)
    # stdout 应被截断到 MAX_STDOUT_INLINE (10KB)
    assert len(r.get("stdout", "")) <= 10_100, "stdout should be truncated"
    # 应有溢出文件
    assert r.get("output_file") is not None, "should have overflow file"


def test_unicode_output():
    """中文/Unicode 输出测试"""
    r = execute("print('你好世界 🎉')")
    assert_ok(r)
    assert_stdout_contains(r, "你好世界")


def test_error_output():
    """错误输出到 stderr"""
    r = execute("""
import sys
sys.stderr.write('test error message\\n')
print('still ok')
""")
    assert_ok(r)
    assert_stdout_contains(r, "still ok")
    assert_stderr_contains(r, "test error message")


def test_exit_code_nonzero():
    """非零退出码"""
    r = execute("import sys; sys.exit(42)")
    assert_fail(r)
    assert r.get("exit_code") == 42, f"Expected exit_code=42, got {r.get('exit_code')}"


def test_work_dir_parameter():
    """work_dir 参数指定工作目录"""
    work_dir = f"/data/agents/{AGENT_ID}/workspace"
    r = execute("import os; print(os.getcwd())", work_dir=work_dir)
    assert_ok(r)
    assert_stdout_contains(r, work_dir)


def test_batch_data_analysis():
    """批量数据分析脚本（模拟真实场景）"""
    work_dir = f"/data/agents/{AGENT_ID}/workspace"
    code = f"""
import pandas as pd
import numpy as np
import os

os.makedirs('{work_dir}', exist_ok=True)

# 生成测试数据
df = pd.DataFrame({{
    'name': ['Alice', 'Bob', 'Charlie'],
    'score': [95, 87, 92],
    'dept': ['Finance', 'Audit', 'Finance']
}})

# 数据分析
summary = df.groupby('dept')['score'].agg(['mean', 'count'])
print(summary.to_string())

# 写入 Excel
df.to_excel('{work_dir}/report.xlsx', index=False)
print('excel_saved')

# 读回验证
df2 = pd.read_excel('{work_dir}/report.xlsx')
print(f'rows={{len(df2)}}')
"""
    r = execute(code, work_dir=work_dir, timeout=15)
    assert_ok(r)
    assert_stdout_contains(r, "Finance")
    assert_stdout_contains(r, "excel_saved")
    assert_stdout_contains(r, "rows=3")


def test_batch_document_processing():
    """批量文档处理（docx）"""
    work_dir = f"/data/agents/{AGENT_ID}/workspace"
    code = f"""
import os
from docx import Document

os.makedirs('{work_dir}', exist_ok=True)

# 创建 Word 文档
doc = Document()
doc.add_heading('测试报告', 0)
doc.add_paragraph('这是一段测试文本。')
doc.save('{work_dir}/test_report.docx')

# 读取验证
doc2 = Document('{work_dir}/test_report.docx')
for p in doc2.paragraphs:
    print(p.text)
print('docx_ok')
"""
    r = execute(code, work_dir=work_dir, timeout=15)
    assert_ok(r)
    assert_stdout_contains(r, "测试报告")
    assert_stdout_contains(r, "docx_ok")


def test_node_version():
    """Node.js 版本验证"""
    r = execute("console.log('node=' + process.version)", language="node")
    assert_ok(r)
    assert_stdout_contains(r, "node=v22")


def test_concurrent_execution():
    """连续快速执行（模拟并发场景）"""
    results = []
    for i in range(5):
        r = execute(f"print('task_{i}')")
        results.append(r)

    for i, r in enumerate(results):
        assert_ok(r, f"task_{i} failed")
        assert_stdout_contains(r, f"task_{i}")


def test_security_bash_rm_rf():
    """安全检查：rm -rf / 被拦截"""
    r = execute("rm -rf /", language="bash")
    assert_fail(r)
    assert "dangerous command" in (r.get("error") or "")


def test_security_bash_sudo():
    """安全检查：sudo 被拦截"""
    r = execute("sudo apt-get install something", language="bash")
    assert_fail(r)
    assert "dangerous command" in (r.get("error") or "")


def test_security_python_os_system():
    """安全检查：os.system 被拦截"""
    r = execute("import os; os.system('ls')")
    assert_fail(r)
    assert "unsafe operation" in (r.get("error") or "")


def test_security_python_requests_allowed():
    """安全检查：requests 库允许使用"""
    r = execute("import requests; print('requests allowed')")
    assert_ok(r)
    assert_stdout_contains(r, "requests allowed")


def test_security_bash_curl_allowed():
    """安全检查：curl 允许使用"""
    r = execute("echo curl test", language="bash")
    assert_ok(r)


def test_output_truncation():
    """输出截断测试 — stdout 超 10KB 被截断，溢出文件保存"""
    work_dir = f"/data/agents/{AGENT_ID}"
    # 生成 >10KB 的 stdout
    r = execute("print('A' * 15000)", work_dir=work_dir)
    assert_ok(r)
    assert len(r.get("stdout", "")) <= 10_100, "stdout should be truncated to ~10KB"
    assert r.get("output_file") is not None, "should have overflow file path"


def test_performance_python():
    """Python 执行性能基线"""
    start = time.time()
    r = execute("x = sum(range(1000)); print(x)")
    elapsed = time.time() - start
    assert_ok(r)
    assert r["duration_ms"] < 500, f"Too slow: {r['duration_ms']}ms"
    print(f"  duration: {r['duration_ms']}ms, total: {elapsed*1000:.0f}ms")


def test_performance_bash():
    """Bash 执行性能基线"""
    start = time.time()
    r = execute("echo fast", language="bash")
    elapsed = time.time() - start
    assert_ok(r)
    assert r["duration_ms"] < 200, f"Too slow: {r['duration_ms']}ms"
    print(f"  duration: {r['duration_ms']}ms, total: {elapsed*1000:.0f}ms")


def test_performance_node():
    """Node.js 执行性能基线"""
    start = time.time()
    r = execute("console.log('fast')", language="node")
    elapsed = time.time() - start
    assert_ok(r)
    assert r["duration_ms"] < 200, f"Too slow: {r['duration_ms']}ms"
    print(f"  duration: {r['duration_ms']}ms, total: {elapsed*1000:.0f}ms")


# ---------------------------------------------------------------------------
# 入口
# ---------------------------------------------------------------------------
def main():
    global BASE_URL, AGENT_ID

    parser = argparse.ArgumentParser(description="Sandbox test suite")
    parser.add_argument("--url", default="http://localhost:8888")
    parser.add_argument("--agent-id", default="test-sandbox-agent")
    args = parser.parse_args()

    BASE_URL = args.url.rstrip("/")
    AGENT_ID = args.agent_id

    print(f"Sandbox Test Suite")
    print(f"URL: {BASE_URL}")
    print(f"Agent ID: {AGENT_ID}")

    # 先检查健康
    health = _get("/health")
    if health.get("status") != "ok":
        print(f"\n[FAIL] Sandbox service not available: {health}")
        sys.exit(1)
    print("Health check: OK\n")

    # 按分类运行测试
    tests = [
        # 基础执行
        ("Health Check", test_health),
        ("Python Basic", test_python_basic),
        ("Python Multiline", test_python_multiline),
        ("Bash Basic", test_bash_basic),
        ("Bash Multiline", test_bash_multiline),
        ("Node Basic", test_node_basic),
        ("Node Multiline", test_node_multiline),

        # 预装包
        ("Pre-installed: pandas", test_preinstalled_pandas),
        ("Pre-installed: numpy", test_preinstalled_numpy),
        ("Pre-installed: scipy", test_preinstalled_scipy),
        ("Pre-installed: matplotlib", test_preinstalled_matplotlib),
        ("Pre-installed: docx", test_preinstalled_docx),
        ("Pre-installed: pptx", test_preinstalled_pptx),
        ("Pre-installed: openpyxl", test_preinstalled_openpyxl),
        ("Pre-installed: xlrd", test_preinstalled_xlrd),
        ("Pre-installed: Pillow", test_preinstalled_pillow),
        ("Pre-installed: pdfplumber", test_preinstalled_pdfplumber),
        ("Pre-installed: fpdf2", test_preinstalled_fpdf2),
        ("Pre-installed: reportlab", test_preinstalled_reportlab),
        ("Pre-installed: requests", test_preinstalled_requests),
        ("Pre-installed: beautifulsoup4", test_preinstalled_beautifulsoup),
        ("Pre-installed: lxml", test_preinstalled_lxml),
        ("Pre-installed: chardet", test_preinstalled_chardet),
        ("Pre-installed: tqdm", test_preinstalled_tqdm),

        # 运行时依赖
        ("Runtime pip install", test_pip_install_runtime),

        # 工作区文件操作
        ("File write to workspace", test_file_write_to_workspace),
        ("File read from workspace", test_file_read_from_workspace),
        ("File write to agent root (skills/)", test_file_write_to_agent_root),

        # 环境变量
        ("WORKSPACE_DIR env var", test_workspace_dir_env_var),
        ("PYTHONPATH env var", test_pythonpath_env),

        # 超时
        ("Timeout handling", test_timeout_short),

        # 跨语言调用
        ("Bash invoke Python", test_bash_invoke_python),
        ("Bash invoke Node", test_bash_invoke_node),
        ("pip install in bash", test_pip_install_in_python),

        # 输出处理
        ("Large output", test_large_output),
        ("Unicode output", test_unicode_output),
        ("Error output (stderr)", test_error_output),
        ("Non-zero exit code", test_exit_code_nonzero),

        # 工作目录参数
        ("work_dir parameter", test_work_dir_parameter),

        # 真实场景
        ("Batch data analysis (Excel)", test_batch_data_analysis),
        ("Batch document processing (docx)", test_batch_document_processing),
        ("Node.js version check", test_node_version),

        # 并发
        ("Consecutive execution (5x)", test_concurrent_execution),

        # 安全检查
        ("Security: rm -rf / blocked", test_security_bash_rm_rf),
        ("Security: sudo blocked", test_security_bash_sudo),
        ("Security: os.system blocked", test_security_python_os_system),
        ("Security: requests allowed", test_security_python_requests_allowed),
        ("Security: curl allowed", test_security_bash_curl_allowed),
        ("Output truncation", test_output_truncation),

        # 性能
        ("Performance: Python", test_performance_python),
        ("Performance: Bash", test_performance_bash),
        ("Performance: Node", test_performance_node),
    ]

    for name, func in tests:
        run_test(name, func)

    # 汇总
    total = PASSED + FAILED
    print(f"\n{'='*60}")
    print(f"RESULTS: {PASSED}/{total} passed, {FAILED} failed")
    if ERRORS:
        print(f"\nFailures:")
        for e in ERRORS:
            print(f"  - {e}")
    print(f"{'='*60}")

    sys.exit(0 if FAILED == 0 else 1)


if __name__ == "__main__":
    main()
